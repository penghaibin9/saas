from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).with_name("E2E学生A_智能分拣系统设计与实现_初稿.docx")
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_GRAY = "F2F4F7"
TEXT = "1F2937"
MUTED = "667085"


def set_run_font(run, size=None, bold=None, color=None, east_asia="Microsoft YaHei"):
    run.font.name = "Calibri"
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_twips):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color="D0D5DD", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    field_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), "Calibri")
    rfonts.set(qn("w:hAnsi"), "Calibri")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "18")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), MUTED)
    rpr.extend([rfonts, sz, color])
    field_run.extend([rpr, begin, instr, separate, value, end])
    paragraph._p.append(field_run)
    tail = paragraph.add_run(" 页")
    set_run_font(tail, size=9, color=MUTED)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    normal.paragraph_format.widow_control = True

    specs = {
        "Title": (24, DARK_BLUE, 0, 14),
        "Subtitle": (12, MUTED, 0, 16),
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in specs.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = name != "Subtitle"
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_body(doc, text, first_line=True):
    p = doc.add_paragraph()
    if first_line:
        p.paragraph_format.first_line_indent = Inches(0.29)
    run = p.add_run(text)
    set_run_font(run, size=11, color=TEXT)
    return p


def add_label_value(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(label)
    set_run_font(r1, size=10, bold=True, color=DARK_BLUE)
    r2 = p.add_run(value)
    set_run_font(r2, size=10, color=TEXT)
    return p


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)

    props = doc.core_properties
    props.title = "智能分拣系统设计与实现（示例初稿）"
    props.subject = "毕业设计统一批阅中心测试材料"
    props.author = "E2E学生A"
    props.keywords = "毕业设计, 智能分拣, 初稿, 评阅测试"

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = header.add_run("毕业设计初稿 · 统一批阅中心测试材料")
    set_run_font(hr, size=9, color=MUTED)
    add_page_field(section.footer.paragraphs[0])

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(10)
    kr = kicker.add_run("毕业设计论文（初稿）")
    set_run_font(kr, size=10, bold=True, color=BLUE)

    title = doc.add_paragraph(style="Title")
    title.add_run("智能分拣系统设计与实现")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("——基于视觉识别与可编程控制的教学型原型")

    doc.add_paragraph()
    add_label_value(doc, "学生：", "E2E学生A")
    add_label_value(doc, "课题：", "E2E课题A-智能分拣系统-20260722-230844")
    add_label_value(doc, "材料类型：", "成果初稿（THESIS_DRAFT）")
    add_label_value(doc, "批次：", "E2E-毕业设计全流程-20260722-230844（2026届）")
    add_label_value(doc, "提交日期：", "2026年8月24日")

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(8)
    note.paragraph_format.space_after = Pt(12)
    note.paragraph_format.left_indent = Inches(0.2)
    note.paragraph_format.right_indent = Inches(0.2)
    nr = note.add_run("摘要预览：本研究面向职业院校实训场景，设计一套由工业相机、传送机构和PLC组成的教学型智能分拣系统，并对识别准确率与节拍进行测试。")
    set_run_font(nr, size=11, color=DARK_BLUE)

    doc.add_page_break()

    doc.add_heading("摘  要", level=1)
    add_body(doc, "针对传统人工分拣效率低、重复劳动强度大等问题，本文设计并实现了一套基于机器视觉的智能分拣系统。系统以工业相机采集工件图像，通过改进的识别算法完成颜色、形状与缺陷判断，再由PLC控制气缸完成分流。样机测试结果表明，系统综合识别准确率达到98.7%，平均处理时间为0.86秒，能够满足中小型实训生产线的需求。", first_line=True)
    keywords = doc.add_paragraph()
    keywords.paragraph_format.space_before = Pt(8)
    k1 = keywords.add_run("关键词：")
    set_run_font(k1, size=11, bold=True, color=DARK_BLUE)
    k2 = keywords.add_run("机器视觉；智能分拣；PLC；职业教育")
    set_run_font(k2, size=11, color=TEXT)

    doc.add_heading("1 研究背景与目标", level=1)
    add_body(doc, "制造企业对柔性化、数字化生产的需求持续提高，智能分拣已经成为典型的自动化应用。职业院校现有实训设备多以单一动作演示为主，缺少视觉识别、控制联动和数据追踪的完整链路。因此，本课题拟构建一套成本可控、结构直观且便于二次开发的教学原型。")
    add_body(doc, "本课题的目标是完成机械结构、电气控制、识别程序和人机界面的集成，并通过连续运行试验验证系统的稳定性。相关行业资料显示，视觉分拣可显著降低漏检率[1]，但本文未进一步说明资料来源、检索范围和对比条件。")

    doc.add_heading("2 总体方案设计", level=1)
    doc.add_heading("2.1 系统组成", level=2)
    add_body(doc, "系统主要由上料模块、传送带、图像采集单元、控制柜、执行气缸和触摸屏构成。工件进入拍摄区域后，光电传感器触发相机采图；工控机完成识别并把类别编码发送给PLC；PLC根据编码和延时参数驱动对应气缸。")
    add_body(doc, "为降低系统成本，原型同时采用24 V直流电源和220 V交流电源。电气设计部分目前只绘制了主回路示意，尚未给出接地、急停回路与气压异常保护的具体选型依据。")

    doc.add_heading("2.2 识别方法", level=2)
    add_body(doc, "图像处理流程包括灰度化、阈值分割、轮廓提取与特征分类。阈值参数由现场调试确定，训练集包含1200张样本图像，其中合格件800张、缺陷件300张。模型在全部样本上训练50轮后，选择准确率最高的一轮用于测试。")
    add_body(doc, "由于初稿阶段尚未划分独立验证集，测试数据与训练数据存在重叠。文中提出的“改进算法”主要调整了阈值和轮廓面积范围，尚未给出与基础方法的消融对比或计算复杂度分析。")

    doc.add_heading("3 样机测试与结果", level=1)
    add_body(doc, "样机在实验室环境连续运行约2小时，传送带速度设为固定档位。测试从三类常见工件中随机抽取样本，记录识别结果、分拣动作与处理时间。由于现场照度未使用照度计测量，下表中的“强光”和“弱光”由操作人员主观判断。")

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["测试场景", "样本量", "识别准确率", "平均耗时", "备注"]
    widths = [2200, 1200, 1500, 1500, 2960]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_run_font(r, size=10, bold=True, color=DARK_BLUE)
    rows = [
        ("正常照明", "150", "96.0%", "0.82 s", "误识别6件"),
        ("强光干扰", "100", "91.0%", "0.94 s", "表面反光明显"),
        ("弱光环境", "80", "92.5%", "0.91 s", "边缘特征缺失"),
        ("合计/平均", "350", "93.8%", "0.89 s", "连续测试约2小时"),
    ]
    for values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx in (1, 2, 3) else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(value)
            set_run_font(r, size=10, color=TEXT)
    set_table_geometry(table, widths)
    set_table_borders(table)

    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(10)
    cr = caption.add_run("表1 不同照明条件下的识别结果")
    set_run_font(cr, size=9, color=MUTED)

    add_body(doc, "测试结果表明，系统总体准确率满足预期，且所有场景均达到95%以上。后续拟增加光源遮罩和自适应曝光，以提高复杂环境下的识别能力。")

    doc.add_heading("4 结论与不足", level=1)
    add_body(doc, "本文完成了智能分拣原型的总体设计与初步测试，实现了工件检测、类别判断和自动分流。受时间和条件限制，当前测试规模较小，尚未覆盖不同传送速度、工件间距和长时间连续运行等工况；下一阶段将补充独立测试集、完善安全回路，并开展对照实验。")

    doc.add_heading("参考文献", level=1)
    refs = [
        "[1] 张某某. 面向智能制造的机器视觉分拣技术研究[J]. 自动化应用, 2024(6): 12-16.",
        "[2] SIEMENS. S7-1200 可编程控制器系统手册[Z]. 2023.",
        "[3] 王某. 工业视觉检测系统开发[M]. 北京: 机械工业出版社, 2022.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.left_indent = Inches(0.25)
        r = p.add_run(ref)
        set_run_font(r, size=10, color=TEXT)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
